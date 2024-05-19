#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed May 15 16:22:54 2024

@author: ashutoshgoenka
"""

import streamlit as st
from PIL import Image
import PIL
import os
import zipfile
from zipfile import ZipFile, ZIP_DEFLATED
import pathlib
import shutil
import docx
import docxtpl
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT


try:
    shutil.rmtree("images_comp")
except:
    pass

name_index_dict= {}

st.title("Resize Images")
# st.write('My first app Hello *world!*')
up_files = st.file_uploader("Upload Image Files", type = ["png", "jpeg", "jpg"] ,accept_multiple_files=True)
# st.write(up_files)

def resize(img, new_width):
    width, height  = img.size
    ratio = height/width
    new_height = int(ratio*new_width)
    resized_image = img.resize((new_width, new_height), resample=PIL.Image.LANCZOS)
    return resized_image


def updateTable():
    # global up_files
    global folder
    document = Document("Table_Word.docx")
    table = document.add_table(rows = 7 , cols = 3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Item'     
    # hdr_cells[1].text = 'quantity'
    document.save("Table_Word.docx")
    counter = 0
    counter_cols = 0
    for file in folder.iterdir():
        cell = table.rows[counter].cells[counter_cols]
        cell._element.clear_content()
        picture = cell.add_paragraph().add_run().add_picture('images_comp/'+file.name, width=Inches(2.6), height=Inches(2.6))
        cell = table.rows[counter+1].cells[counter_cols]
        name = os.path.splitext(file.name)[0]
        cell.text = name
        if counter_cols<2:
            counter_cols = counter_cols + 1
        else:
            counter_cols = 0
            counter = counter+2
    document.save("Table_Word.docx")
    
                
    

title = st.text_input("Image Number to Start with", 1)
st.write("The image numbering will start from: ", title)

try:
    os.mkdir("images_comp")
except:
    pass

name_list = []

for i in range(len(up_files)):
    name_list.append("Image "+str(int(title)+i))



st.write(len(up_files))


for file in up_files:
    try:
        a = name_index_dict[file.name]
    except:
        name_index_dict[file.name] = 0

    
    # files = os.listdir("images")
    extensions = ["jpg", "jpeg", "png", "gif", "webp"]
    im = Image.open(file)
    ext = file.name.split(".")[-1]
    # st.write(file)
    st.image(file, width=250)
    name = os.path.splitext(file.name)[0]
    option =   st.selectbox(
            "File Name",
            tuple([name] + name_list),
            index=name_index_dict[file.name],
            )

    st.write("You selected:", option)
    list_temp = [name] + name_list
    position = list_temp.index(option)

    name_index_dict[name] = position
    # st.write(im.size)
    # im_resized = resize(im, 1000)
    # st.write(im_resized.size)
    im_width, im_height = im.size 
    st.write(im_width, im_height)
    size_to_scale = min(im_width,im_height)
    st.write(size_to_scale)
    # box = (size_to_scale, size_to_scale, size_to_scale, size_to_scale)
    box=((im_width-size_to_scale)/2,((im_height-size_to_scale)/2),(im_width+size_to_scale)/2,((im_height+size_to_scale)/2))
    im_resized = im.crop(box)
    # st.image(im_resized)    
    im_resized.save("images_comp/"+option+"."+ext)
    
    

zip_path = "images_compressed.zip"
directory_to_zip = "images_comp"
folder = pathlib.Path(directory_to_zip)
# st.write(folder)

#Create a document with Landscape and saved
document = Document()
section = document.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height
document.save("Table_Word.docx")
# Document Created


with ZipFile(zip_path, 'w', ZIP_DEFLATED) as zip:
    for file in folder.iterdir():
        zip.write(file, arcname=file.name)
        
with open("images_compressed.zip", "rb") as fp:
    btn = st.download_button(
        label="Download ZIP",
        data=fp,
        file_name="images_compressed.zip",
        mime="application/zip"
    )

updateTable()

with open("Table_Word.docx", "rb") as fp:

    btn_1 = st.download_button(
            label="Word File with Table",
            data=fp,
            file_name="Table_Word_docx",
            mime="docx"
            )
    
    # st.write(btn_1)
    
    # if btn_1:
    #     st.write("Running Update Function")
    #     updateTable(up_files)

os.remove(zip_path)
shutil.rmtree("images_comp")

# st.download_button("Download Images", file_name="bali.jpeg")
    
    # for file in files:
    # ext = im.name.split(".")[-1]
    # if ext in extensions:
    #     # im = Image.open("images/"+file)
        
        
    #     im_resized = resize(im, 400)
    #     filepath = "images/"+file+".jpg"
    #     im_resized.save(filepath)
        
